"""텍스트·docx·hwp에서 대본 문자열 추출."""

from __future__ import annotations

from contextlib import closing
from io import StringIO
from pathlib import Path


def extract_text_from_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in (".txt", ".md"):
        return _read_plain_text(path)
    if suffix == ".doc":
        raise ValueError(
            "구형 Word(.doc)은 지원하지 않습니다. Word에서 .docx로 저장하거나 "
            "Google 문서는 '파일 → 다운로드 → Microsoft Word(.docx)'로 내보내 주세요."
        )
    if suffix == ".docx":
        return _text_from_docx(path)
    if suffix == ".hwp":
        return _text_from_hwp(path)
    raise ValueError(f"지원하지 않는 형식입니다: {suffix}")


def _read_plain_text(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _text_from_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" ".join(cells))
    return "\n".join(parts)


def _text_from_hwp(path: Path) -> str:
    from hwp5.hwp5txt import TextTransform
    from hwp5.xmlmodel import Hwp5File

    text_transform = TextTransform()
    transform = text_transform.transform_hwp5_to_text
    buf = StringIO()
    with closing(Hwp5File(str(path))) as hwp5file:
        transform(hwp5file, buf)
    return buf.getvalue()
